import docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_hyperlink(paragraph, url, text, color, underline):
    """
    在段落中添加超链接
    :param paragraph: 段落对象
    :param url: 超链接地址
    :param text: 超链接显示的文本
    :param color: 超链接颜色
    :param underline: 是否下划线
    :return: 超链接对象
    """
    # 添加超链接
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # 创建超链接元素
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # 创建运行元素
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # 设置颜色
    if not color is None:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    # 设置下划线
    u = OxmlElement('w:u')
    u.set(qn('w:val'), underline)
    rPr.append(u)

    # 添加属性到运行元素
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # 将超链接添加到段落
    paragraph._p.append(hyperlink)

    return hyperlink


# 创建一个新的Word文档
doc = Document()

# 添加一个段落
paragraph = doc.add_paragraph()

# 在段落中添加文本
paragraph.add_run('这是一个超链接的示例：')

# 在段落中添加超链接
hyperlink = add_hyperlink(paragraph, "https://www.google.com", "点击这里访问Google", "0000FF", "single")

# 保存文档
doc.save('hyperlink_example.docx')
